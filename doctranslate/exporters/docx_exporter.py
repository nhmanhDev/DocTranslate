import shutil
import os
from docx import Document
from doctranslate.schema import UnifiedDocument

def export_docx(
    template_path: str,
    output_path: str,
    translated_doc: UnifiedDocument,
    mode: str = "mono"  # "mono" or "dual"
):
    """
    Export the translated document to DOCX.
    Uses the original file as a template to preserve all layouts, styles, and media.
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template DOCX file not found: {template_path}")

    # Copy the template to output path to preserve everything
    shutil.copy2(template_path, output_path)
    doc = Document(output_path)

    # Build a lookup dictionary of blocks by ID
    blocks_by_id = {}
    
    def collect_blocks(blocks):
        for block in blocks:
            blocks_by_id[block.id] = block
            if block.table_rows:
                for row in block.table_rows:
                    for cell in row:
                        collect_blocks(cell)
                        
    collect_blocks(translated_doc.blocks)

    block_counter = 0

    def apply_translation_to_paragraph(p, block_id):
        block = blocks_by_id.get(block_id)
        if not block or not block.translated_content:
            return

        translated_text = block.translated_content

        if mode == "mono":
            # Replaces the text of the paragraph while trying to preserve style.
            # We set the text on the first run and clear text on subsequent runs.
            if p.runs:
                p.runs[0].text = translated_text
                for run in p.runs[1:]:
                    run.text = ""
            else:
                p.text = translated_text
        elif mode == "dual":
            # In dual mode, we append the translation as a new line or a new paragraph.
            # Inserting a new paragraph below the current one is cleaner.
            # We add a run in italic/colored text.
            # Since python-docx doesn't easily let us insert a paragraph after a specific paragraph,
            # we can append the translation inside the same paragraph on a new line.
            # This preserves original layout and ordering perfectly!
            p.add_run("\n")
            trans_run = p.add_run(translated_text)
            trans_run.italic = True
            # Let's try to set a slightly grey color if possible to make it look premium
            try:
                from docx.shared import RGBColor
                trans_run.font.color.rgb = RGBColor(128, 128, 128)
            except Exception:
                pass

    # Traverse body elements in the document and apply translation
    for element in doc.element.body:
        tag = element.tag.split("}")[-1]
        if tag == "p":
            p_obj = None
            for p in doc.paragraphs:
                if p._element is element:
                    p_obj = p
                    break
            if p_obj:
                block_counter += 1
                apply_translation_to_paragraph(p_obj, f"p_{block_counter}")
        elif tag == "tbl":
            t_obj = None
            for t in doc.tables:
                if t._element is element:
                    t_obj = t
                    break
            if t_obj:
                block_counter += 1
                cell_p_counter = 0
                for row in t_obj.rows:
                    for cell in row.cells:
                        for cp in cell.paragraphs:
                            cell_p_counter += 1
                            block_id = f"p_t_{block_counter}_{cell_p_counter}"
                            apply_translation_to_paragraph(cp, block_id)

    doc.save(output_path)

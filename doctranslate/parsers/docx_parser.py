import os
from docx import Document
from doctranslate.schema import UnifiedDocument, DocumentBlock, Span

def parse_docx(file_path: str) -> UnifiedDocument:
    """Parse a DOCX file into a UnifiedDocument model."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"DOCX file not found: {file_path}")

    doc = Document(file_path)
    unified_doc = UnifiedDocument(metadata={"title": os.path.basename(file_path), "format": "docx"})

    block_counter = 0

    def parse_paragraph_to_block(p, block_id) -> DocumentBlock:
        content = p.text
        # Determine block type from style
        style_name = p.style.name
        block_type = "paragraph"
        level = None

        if style_name.startswith("Heading"):
            block_type = "heading"
            try:
                level = int(style_name.replace("Heading", "").strip())
            except ValueError:
                level = 1
        elif style_name.startswith("List"):
            block_type = "list_item"

        spans = []
        for run in p.runs:
            # Extract color if present
            color = None
            if run.font.color and run.font.color.rgb:
                color = str(run.font.color.rgb)

            span = Span(
                text=run.text,
                bold=bool(run.bold),
                italic=bool(run.italic),
                underline=bool(run.underline),
                font_size=run.font.size.pt if run.font.size else None,
                font_name=run.font.name,
                color=color,
            )
            spans.append(span)

        # Set block attributes
        attributes = {
            "style_name": style_name,
            "alignment": str(p.alignment) if p.alignment is not None else None,
        }
        if level is not None:
            attributes["level"] = level

        return DocumentBlock(
            id=block_id,
            type=block_type,
            content=content,
            spans=spans,
            attributes=attributes,
        )

    # DOCX contains paragraphs and tables at the document body level.
    # We can iterate through all child elements to preserve their sequential order.
    # We can use doc.element.body to iterate in order.
    for element in doc.element.body:
        tag = element.tag.split("}")[-1]
        if tag == "p":
            # Find the paragraph object corresponding to this element
            p_obj = None
            for p in doc.paragraphs:
                if p._element is element:
                    p_obj = p
                    break
            if p_obj:
                block_counter += 1
                unified_doc.blocks.append(parse_paragraph_to_block(p_obj, f"p_{block_counter}"))
        elif tag == "tbl":
            # Find the table object
            t_obj = None
            for t in doc.tables:
                if t._element is element:
                    t_obj = t
                    break
            if t_obj:
                block_counter += 1
                table_rows = []
                for row in t_obj.rows:
                    row_cells = []
                    for cell in row.cells:
                        # Parse paragraphs inside each cell
                        cell_blocks = []
                        cell_p_counter = 0
                        for cp in cell.paragraphs:
                            cell_p_counter += 1
                            block = parse_paragraph_to_block(cp, f"p_t_{block_counter}_{cell_p_counter}")
                            cell_blocks.append(block)
                        row_cells.append(cell_blocks)
                    table_rows.append(row_cells)

                # Add table block
                unified_doc.blocks.append(DocumentBlock(
                    id=f"table_{block_counter}",
                    type="table",
                    content="",  # Tables don't have direct single-string content
                    table_rows=table_rows,
                    attributes={"rows": len(t_obj.rows), "cols": len(t_obj.columns)}
                ))

    return unified_doc
